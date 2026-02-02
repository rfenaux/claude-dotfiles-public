"""Word document parser using python-docx."""
from pathlib import Path
from typing import Dict, Any
from . import register_parser

class DocxParser:
    extensions = [".docx"]

    def parse(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("DOCX parsing requires python-docx")

        doc = Document(str(file_path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve some structure with headers
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        text = "#" * level_num + " " + text
                    except ValueError:
                        pass
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Word document."""
        stat = file_path.stat()
        metadata = {
            "filename": file_path.name,
            "extension": file_path.suffix,
            "size_bytes": stat.st_size,
            "modified": stat.st_mtime,
        }

        try:
            from docx import Document
            doc = Document(str(file_path))
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title or file_path.stem,
                "author": core_props.author or "",
                "created": str(core_props.created) if core_props.created else "",
            })
        except ImportError:
            metadata["title"] = file_path.stem

        return metadata

# Register the parser
register_parser(DocxParser())
